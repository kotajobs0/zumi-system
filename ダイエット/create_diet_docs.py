from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), val.get('sz', '4'))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color='1F3864', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_section_title(doc, text, bg_color='1F3864', font_color='FFFFFF'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(16)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(font_color)
    doc.add_paragraph()

def add_menu_table(doc, headers, rows, header_color='2E75B6', row_colors=('EBF3FB', 'FFFFFF')):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string('FFFFFF')

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = row_colors[r_idx % 2]
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx + 1, c_idx)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(9)

    # Column widths
    widths = [Cm(2.5), Cm(9), Cm(2.5)] if len(headers) == 3 else [Cm(4)] * len(headers)
    for i, col in enumerate(table.columns):
        if i < len(widths):
            for cell in col.cells:
                cell.width = widths[i]

    doc.add_paragraph()

def add_points_table(doc, points, color='1F3864'):
    p = doc.add_paragraph()
    run = p.add_run('■ ポイント')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)

    for point in points:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(point)
        run.font.size = Pt(10)
    doc.add_paragraph()

def add_rules_table(doc, rules):
    add_section_title(doc, '共通ルール', bg_color='404040')
    table = doc.add_table(rows=len(rules), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (label, val) in enumerate(rules):
        bg = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        set_cell_bg(c0, 'D9D9D9')
        set_cell_bg(c1, bg)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9)
        c0.width = Cm(4)
        c1.width = Cm(12)

# ===================== 女性版 =====================
doc_f = Document()
doc_f.core_properties.author = 'Zumi System'

section = doc_f.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

add_heading(doc_f, 'ダイエットメニュー【女性版】', size=20, color='C0004B')
add_heading(doc_f, '1日目標カロリー：約1,400kcal', size=12, color='555555', bold=False)
doc_f.add_paragraph()

add_section_title(doc_f, '1日のメニュー', bg_color='C0004B')

menu_headers = ['食事', 'メニュー', 'カロリー']
menu_rows_f = [
    ['朝食', 'ギリシャヨーグルト + ベリー類 + 全粒粉トースト1枚', '約350kcal'],
    ['昼食', '鶏むね肉のサラダ（レモンドレッシング）+ 玄米100g + 味噌汁', '約450kcal'],
    ['間食', 'ナッツ20g + 無糖緑茶', '約120kcal'],
    ['夕食', '鮭の蒸し焼き + 野菜炒め（ブロッコリー・パプリカ）+ 豆腐', '約480kcal'],
]
add_menu_table(doc_f, menu_headers, menu_rows_f, header_color='C0004B', row_colors=('FDE8F0', 'FFFFFF'))

points_f = [
    '鉄分・カルシウムを意識して摂取',
    '水分は1日1.5〜2L を目標に',
    '間食はナッツや小魚など良質な脂質を選ぶ',
    'ホルモンバランスに配慮し無理な制限は避ける',
]
add_points_table(doc_f, points_f, color='C0004B')

rules = [
    ['禁止・制限', '揚げ物、砂糖入り飲料、アルコール、白米（玄米に変更）'],
    ['推奨調理法', '蒸す・焼く・茹でる'],
    ['食事の間隔', '4〜5時間おきに規則正しく'],
    ['就寝前', '就寝2時間前は食事を避ける'],
]
add_rules_table(doc_f, rules)

doc_f.save(r'c:\.github\zumi sysytem\ダイエット\ダイエットメニュー_女性版.docx')
print('女性版 作成完了')

# ===================== 男性版 =====================
doc_m = Document()
doc_m.core_properties.author = 'Zumi System'

section_m = doc_m.sections[0]
section_m.page_width = Cm(21)
section_m.page_height = Cm(29.7)
section_m.left_margin = Cm(2)
section_m.right_margin = Cm(2)
section_m.top_margin = Cm(2)
section_m.bottom_margin = Cm(2)

add_heading(doc_m, 'ダイエットメニュー【男性版】', size=20, color='1F3864')
add_heading(doc_m, '1日目標カロリー：約1,800kcal', size=12, color='555555', bold=False)
doc_m.add_paragraph()

add_section_title(doc_m, '1日のメニュー', bg_color='1F3864')

menu_rows_m = [
    ['朝食', '卵2個（スクランブル）+ 全粒粉トースト2枚 + バナナ1本', '約500kcal'],
    ['昼食', '鶏むね肉200g + 玄米150g + 野菜たっぷり味噌汁 + ほうれん草ソテー', '約600kcal'],
    ['間食', 'プロテインシェイク または ゆで卵2個', '約150kcal'],
    ['夕食', '赤身牛肉または鶏もも肉（皮なし）150g + 蒸し野菜 + 豆腐 + わかめスープ', '約550kcal'],
]
add_menu_table(doc_m, menu_headers, menu_rows_m, header_color='1F3864', row_colors=('EBF3FB', 'FFFFFF'))

points_m = [
    'たんぱく質を体重×1.5〜2g/日を目標に摂取',
    '筋トレと組み合わせると効果的',
    '夜の炭水化物は控えめに',
    '水分は1日2L以上を目標に',
]
add_points_table(doc_m, points_m, color='1F3864')
add_rules_table(doc_m, rules)

doc_m.save(r'c:\.github\zumi sysytem\ダイエット\ダイエットメニュー_男性版.docx')
print('男性版 作成完了')
