# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── helpers ───────────────────────────────────────────────────────────────────

def cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=9, color='000000',
              align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)

def add_title(doc, title, sub, accent):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(accent)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub)
    r2.font.size = Pt(11); r2.font.color.rgb = RGBColor.from_string('666666')
    doc.add_paragraph()

def add_day_block(doc, day_jp, meals, accent, row_colors):
    """meals = list of (meal_name, menu_text, kcal)"""

    # ── day header bar ──
    hdr = doc.add_table(rows=1, cols=1)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.columns[0].width = Cm(17)
    hc = hdr.cell(0, 0)
    cell_bg(hc, accent)
    cell_text(hc, f'  {day_jp}', bold=True, size=12, color='FFFFFF',
              align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── meal table ──
    tbl = doc.add_table(rows=len(meals) + 1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # column widths
    widths = [Cm(2.4), Cm(10.8), Cm(2.8)]
    col_headers = ['食事', 'メニュー', 'カロリー']
    header_bg = _darken(accent)

    for ci, (htext, w) in enumerate(zip(col_headers, widths)):
        hcell = tbl.cell(0, ci)
        cell_bg(hcell, header_bg)
        cell_text(hcell, htext, bold=True, size=9, color='FFFFFF',
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        for row in tbl.rows:
            row.cells[ci].width = w

    for ri, (name, menu, kcal) in enumerate(meals):
        bg = row_colors[ri % 2]
        c0 = tbl.cell(ri + 1, 0)
        c1 = tbl.cell(ri + 1, 1)
        c2 = tbl.cell(ri + 1, 2)
        cell_bg(c0, bg); cell_bg(c1, bg); cell_bg(c2, bg)
        cell_text(c0, name, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(c1, menu, size=9)
        cell_text(c2, kcal, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # small spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

def _darken(hex6):
    """Darken a hex color by ~20%."""
    r, g, b = int(hex6[0:2],16), int(hex6[2:4],16), int(hex6[4:6],16)
    r = max(0, int(r * 0.75))
    g = max(0, int(g * 0.75))
    b = max(0, int(b * 0.75))
    return f'{r:02X}{g:02X}{b:02X}'

def add_rules(doc, rules, accent):
    p = doc.add_paragraph()
    r = p.add_run('■ 共通ルール')
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(accent)
    tbl = doc.add_table(rows=len(rules), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    row_colors = ['F2F2F2', 'FFFFFF']
    for i, (label, val) in enumerate(rules):
        c0 = tbl.cell(i, 0); c1 = tbl.cell(i, 1)
        cell_bg(c0, 'D9D9D9'); cell_bg(c1, row_colors[i % 2])
        cell_text(c0, label, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(c1, val, size=9)
        c0.width = Cm(3.5); c1.width = Cm(13.5)

def make_doc(accent, title, sub, week_data, row_colors, rules, outpath):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2)
    sec.top_margin  = sec.bottom_margin = Cm(2)

    add_title(doc, title, sub, accent)

    days_jp = ['月曜日', '火曜日', '水曜日', '木曜日', '金曜日', '土曜日', '日曜日']
    for day_jp, meals in zip(days_jp, week_data):
        add_day_block(doc, day_jp, meals, accent, row_colors)

    doc.add_paragraph()
    add_rules(doc, rules, accent)
    doc.save(outpath)
    print(f'保存: {outpath}')

# ─── data ──────────────────────────────────────────────────────────────────────

RULES = [
    ['禁止・制限', '揚げ物、砂糖入り飲料、アルコール、白米（→玄米に変更）'],
    ['推奨調理法', '蒸す・焼く・茹でる（油は控えめに）'],
    ['食事の間隔', '4〜5時間おきに規則正しく'],
    ['就寝前',     '就寝2時間前は食事を避ける'],
    ['水分',       '1日1.5〜2L（女性）/ 2L以上（男性）'],
]

# ── 女性版 1,400 kcal ─────────────────────────────────────────────────────────
FEMALE_WEEK = [
    # 月
    [('朝食', 'ギリシャヨーグルト + ベリー類 + 全粒粉トースト1枚',              '約350kcal'),
     ('昼食', '鶏むね肉のサラダ（レモンドレッシング）+ 玄米100g + 味噌汁',      '約450kcal'),
     ('間食', 'ナッツ20g + 無糖緑茶',                                            '約120kcal'),
     ('夕食', '鮭の蒸し焼き + 野菜炒め（ブロッコリー・パプリカ）+ 豆腐',         '約480kcal')],
    # 火
    [('朝食', 'オートミール80g + バナナ1/2本 + 無糖豆乳150ml',                   '約320kcal'),
     ('昼食', '豆腐と野菜のスープ煮 + 玄米おにぎり1個',                           '約420kcal'),
     ('間食', 'りんご1/2個 + チーズ1枚',                                          '約130kcal'),
     ('夕食', '蒸し鶏 + ほうれん草のごま和え + 豆腐味噌汁',                       '約530kcal')],
    # 水
    [('朝食', 'ゆで卵1個 + 全粒粉トースト1枚 + オレンジ1/2個',                   '約330kcal'),
     ('昼食', 'サバの塩焼き + 玄米100g + きのこ汁',                               '約460kcal'),
     ('間食', '無糖ヨーグルト100g + はちみつ少量',                                '約100kcal'),
     ('夕食', '豚もも肉の蒸し物 + 蒸し野菜 + わかめスープ',                       '約510kcal')],
    # 木
    [('朝食', 'スムージー（ほうれん草・バナナ・豆乳）+ ゆで卵1個',                '約350kcal'),
     ('昼食', 'ツナと野菜のサラダ + 玄米100g + 味噌汁',                           '約440kcal'),
     ('間食', 'ナッツ15g + 無糖コーヒー',                                          '約 90kcal'),
     ('夕食', '鶏むね肉のグリル + 温野菜 + 豆腐スープ',                           '約520kcal')],
    # 金
    [('朝食', '全粒粉パン2枚 + アボカド1/4個 + ゆで卵1個',                        '約370kcal'),
     ('昼食', '鮭フレーク + 玄米おにぎり + 野菜スープ',                            '約430kcal'),
     ('間食', 'キウイ1個 + 無糖ヨーグルト100g',                                   '約110kcal'),
     ('夕食', 'えびと野菜の蒸し物 + 豆腐 + わかめ味噌汁',                         '約490kcal')],
    # 土
    [('朝食', 'オートミールパンケーキ + 無糖ヨーグルト',                           '約340kcal'),
     ('昼食', '鶏むね肉のレモン蒸し + 玄米100g + きのこ汁',                       '約450kcal'),
     ('間食', '煮干し15g + 緑茶',                                                  '約 60kcal'),
     ('夕食', '白身魚の蒸し焼き + ブロッコリー + 豆腐スープ',                     '約550kcal')],
    # 日
    [('朝食', 'ギリシャヨーグルト + ベリー類 + グラノーラ20g',                    '約360kcal'),
     ('昼食', '豆腐丼（玄米100g + 絹豆腐 + 野菜）+ 味噌汁',                      '約420kcal'),
     ('間食', 'ナッツ20g + 無糖紅茶',                                              '約120kcal'),
     ('夕食', '鶏のグリル + サラダ（ノンオイルドレッシング）+ 野菜スープ',         '約500kcal')],
]

# ── 男性版 1,800 kcal ─────────────────────────────────────────────────────────
MALE_WEEK = [
    # 月
    [('朝食', '卵2個（スクランブル）+ 全粒粉トースト2枚 + バナナ1本',              '約500kcal'),
     ('昼食', '鶏むね肉200g + 玄米150g + 野菜たっぷり味噌汁 + ほうれん草ソテー', '約600kcal'),
     ('間食', 'プロテインシェイク または ゆで卵2個',                               '約150kcal'),
     ('夕食', '赤身牛肉150g + 蒸し野菜 + 豆腐 + わかめスープ',                    '約550kcal')],
    # 火
    [('朝食', 'オートミール150g + バナナ1本 + 無糖豆乳200ml',                      '約480kcal'),
     ('昼食', 'サバの塩焼き + 玄米150g + 豆腐 + 味噌汁',                          '約580kcal'),
     ('間食', 'ゆで卵2個 + ナッツ20g',                                             '約200kcal'),
     ('夕食', '豚もも肉200g + 蒸し野菜 + わかめスープ + 豆腐',                    '約540kcal')],
    # 水
    [('朝食', '卵3個（スクランブル）+ 全粒粉トースト2枚 + オレンジ1個',            '約520kcal'),
     ('昼食', '鮭150g + 玄米150g + 野菜炒め + 味噌汁',                            '約590kcal'),
     ('間食', 'プロテインシェイク',                                                 '約150kcal'),
     ('夕食', '鶏むね肉200g + 蒸し野菜 + 豆腐スープ',                             '約540kcal')],
    # 木
    [('朝食', '全粒粉パン3枚 + アボカド1/2個 + ゆで卵2個',                         '約510kcal'),
     ('昼食', 'ツナ缶（水煮）+ 玄米150g + 野菜スープ + サラダ',                   '約570kcal'),
     ('間食', 'ゆで卵2個 + 無糖ヨーグルト',                                        '約180kcal'),
     ('夕食', '赤身牛肉150g + ブロッコリー + 豆腐 + わかめスープ',                 '約540kcal')],
    # 金
    [('朝食', 'オートミール + バナナ + ゆで卵2個 + 無糖豆乳',                      '約490kcal'),
     ('昼食', '鶏むね肉200g + 玄米150g + ほうれん草ソテー + 味噌汁',              '約600kcal'),
     ('間食', 'プロテインシェイク + ナッツ20g',                                    '約200kcal'),
     ('夕食', 'えびと野菜の炒め物 + 豆腐 + わかめスープ',                          '約510kcal')],
    # 土
    [('朝食', '卵2個（目玉焼き）+ 全粒粉トースト2枚 + バナナ1本',                  '約500kcal'),
     ('昼食', '鮭200g + 玄米150g + きのこ汁 + 野菜ソテー',                        '約610kcal'),
     ('間食', 'ゆで卵2個',                                                          '約160kcal'),
     ('夕食', '豚もも肉150g + 蒸し野菜 + 豆腐スープ',                             '約530kcal')],
    # 日
    [('朝食', 'オートミールパンケーキ + ゆで卵2個 + バナナ1本',                    '約510kcal'),
     ('昼食', '鶏むね肉200g + 玄米150g + 野菜炒め + 味噌汁',                      '約590kcal'),
     ('間食', 'プロテインシェイク',                                                 '約150kcal'),
     ('夕食', '赤身牛肉150g + 蒸し野菜 + 豆腐 + わかめスープ',                    '約550kcal')],
]

# ─── generate ─────────────────────────────────────────────────────────────────

make_doc(
    accent      = 'C0004B',
    title       = 'ダイエットメニュー【女性版】1週間プラン',
    sub         = '1日目標カロリー：約1,400kcal　／　たんぱく質・鉄分・カルシウムを重視',
    week_data   = FEMALE_WEEK,
    row_colors  = ('FDE8F0', 'FFFFFF'),
    rules       = RULES,
    outpath     = r'c:\.github\zumi sysytem\ダイエット\ダイエットメニュー_女性版_1週間.docx',
)

make_doc(
    accent      = '1F3864',
    title       = 'ダイエットメニュー【男性版】1週間プラン',
    sub         = '1日目標カロリー：約1,800kcal　／　高たんぱく・筋トレ併用を想定',
    week_data   = MALE_WEEK,
    row_colors  = ('EBF3FB', 'FFFFFF'),
    rules       = RULES,
    outpath     = r'c:\.github\zumi sysytem\ダイエット\ダイエットメニュー_男性版_1週間.docx',
)
