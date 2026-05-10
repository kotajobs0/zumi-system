from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ余白
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ヘルパー：段落フォント設定
def set_font(run, size=11, bold=False, color=None, font_name="游ゴシック"):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 日本語フォント設定
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rpr.insert(0, rFonts)

def set_para_font(para, font_name="游ゴシック"):
    ppr = para._p.get_or_add_pPr()
    rpr_default = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rpr_default.append(rFonts)
    ppr.append(rpr_default)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=(31, 73, 125)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True, color=color)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

# ===== タイトル =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("FP3級 合格ロードマップ（〜2026年9月）")
set_font(title_run, size=18, bold=True, color=(31, 73, 125))
title_p.paragraph_format.space_after = Pt(4)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("作成日：2026年5月10日　　目標試験日：2026年9月6日（日）")
set_font(sub_run, size=10, color=(128, 128, 128))
sub_p.paragraph_format.space_after = Pt(10)

doc.add_paragraph()

# ===== 試験概要 =====
add_heading(doc, "■ 試験概要")

table1 = doc.add_table(rows=6, cols=2)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ["項目", "内容"]
rows_data = [
    ("試験日", "2026年9月6日（日）予定"),
    ("科目", "学科（60問）＋ 実技（3科目から選択）"),
    ("合格基準", "学科：36点以上 / 実技：60点以上"),
    ("合格率", "約60〜80%（比較的取りやすい）"),
    ("必要学習時間", "50〜100時間"),
]

header_row = table1.rows[0]
for i, h in enumerate(headers):
    cell = header_row.cells[i]
    shade_cell(cell, "1F497D")
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_font(run, size=10.5, bold=True, color=(255, 255, 255))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for row_idx, (label, value) in enumerate(rows_data):
    row = table1.rows[row_idx + 1]
    if row_idx % 2 == 0:
        shade_cell(row.cells[0], "DCE6F1")
        shade_cell(row.cells[1], "EBF3FB")
    else:
        shade_cell(row.cells[0], "F2F2F2")
        shade_cell(row.cells[1], "FFFFFF")
    r0 = row.cells[0].paragraphs[0].add_run(label)
    set_font(r0, size=10.5, bold=True)
    r1 = row.cells[1].paragraphs[0].add_run(value)
    set_font(r1, size=10.5)

# 列幅
for row in table1.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(11)

doc.add_paragraph()

# ===== 1日の勉強時間 =====
add_heading(doc, "■ 1日の勉強時間の目安")

time_p = doc.add_paragraph()
time_run = time_p.add_run("平日：1〜1.5時間　／　休日：2〜3時間　→　合計約150時間確保")
set_font(time_run, size=11, bold=True, color=(192, 80, 77))

note_p = doc.add_paragraph()
note_run = note_p.add_run("※ 合計150時間を確保することで、余裕をもって合格圏内に入ることができます。")
set_font(note_run, size=10, color=(128, 128, 128))
note_p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

# ===== ロードマップ =====
add_heading(doc, "■ 4ヶ月ロードマップ")

phases = [
    ("Phase 1　インプット期（5月〜6月中旬）", "1F497D", [
        "テキストを1周読む（理解優先・暗記は不要）",
        "6分野を順番に学習：",
        "　① ライフプランニングと資金計画　② リスク管理（保険）　③ 金融資産運用",
        "　④ タックスプランニング（税金）　⑤ 不動産　⑥ 相続・事業承継",
    ]),
    ("Phase 2　アウトプット期（6月中旬〜7月末）", "375623", [
        "過去問を繰り返し解く（5年分が目安）",
        "間違えた問題をノートにまとめる",
        "学科・実技の両方を対策する",
    ]),
    ("Phase 3　仕上げ期（8月〜試験直前）", "843C0C", [
        "模擬試験を本番形式で解く",
        "苦手分野を集中復習する",
        "法改正・最新情報を確認（税率・年金額など）",
    ]),
]

for phase_title, color_hex, items in phases:
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    p = doc.add_paragraph()
    run = p.add_run(f"◆ {phase_title}")
    set_font(run, size=11, bold=True, color=(r, g, b))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    for item in items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent = Cm(0.5)
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(1)
        run = bp.add_run(item)
        set_font(run, size=10.5)

doc.add_paragraph()

# ===== 参考書 =====
add_heading(doc, "■ おすすめ参考書")

add_heading(doc, "テキスト", level=2, color=(31, 73, 125))

table2 = doc.add_table(rows=3, cols=2)
table2.style = 'Table Grid'
books = [
    ("書籍名", "特徴"),
    ("みんなが欲しかった！FPの教科書3級（TAC出版）", "図解多め・初心者に最もわかりやすい。定番中の定番"),
    ("FP技能士3級テキスト＆問題集（ナツメ社）", "テキスト＋問題集が1冊にまとまっており効率的"),
]

for i, (col1, col2) in enumerate(books):
    row = table2.rows[i]
    if i == 0:
        shade_cell(row.cells[0], "4F81BD")
        shade_cell(row.cells[1], "4F81BD")
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5, bold=True, color=(255, 255, 255))
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5, bold=True, color=(255, 255, 255))
    else:
        fill = "DCE6F1" if i % 2 == 1 else "FFFFFF"
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], fill)
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5)
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5)

for row in table2.rows:
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(8)

doc.add_paragraph()
add_heading(doc, "問題集", level=2, color=(31, 73, 125))

table3 = doc.add_table(rows=3, cols=2)
table3.style = 'Table Grid'
mondai = [
    ("書籍名", "特徴"),
    ("みんなが欲しかった！FPの問題集3級（TAC出版）", "テキストと連動・解説が丁寧"),
    ("FP3級 過去問題集（きんざい）", "本番に近い形式・実技対策も充実"),
]

for i, (col1, col2) in enumerate(mondai):
    row = table3.rows[i]
    if i == 0:
        shade_cell(row.cells[0], "4F81BD")
        shade_cell(row.cells[1], "4F81BD")
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5, bold=True, color=(255, 255, 255))
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5, bold=True, color=(255, 255, 255))
    else:
        fill = "DCE6F1" if i % 2 == 1 else "FFFFFF"
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], fill)
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5)
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5)

for row in table3.rows:
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(8)

doc.add_paragraph()
add_heading(doc, "無料サービス（おすすめ）", level=2, color=(31, 73, 125))

free_items = [
    ("FP3級ドットコム", "過去問が無料で解け、解説も詳しい"),
    ("学科試験 過去問道場", "スマホで隙間時間に活用できる"),
]
for name, desc in free_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"・ {name}  —  {desc}")
    set_font(r, size=10.5)

doc.add_paragraph()

# ===== YouTubeチャンネル =====
add_heading(doc, "■ おすすめYouTubeチャンネル")

yt_note = doc.add_paragraph()
note_r = yt_note.add_run("※ 各チャンネル名でYouTube検索してください。")
set_font(note_r, size=9.5, color=(128, 128, 128))

table4 = doc.add_table(rows=5, cols=2)
table4.style = 'Table Grid'
yt_data = [
    ("チャンネル名", "特徴・活用方法"),
    ("ほんださん / 東大式FPチャンネル", "FP学習系で最も人気。図解でわかりやすく、3級〜1級まで対応。再生リストが体系的で学習順に見やすい"),
    ("聞くだけFP", "音声メインで「ながら学習」に最適。通勤・家事の合間に耳で勉強できる"),
    ("資格の学校TAC公式", "大手予備校公式。解説の信頼性が高い"),
    ("フォーサイト公式", "通信講座のサンプル講義が無料で視聴可能"),
]

for i, (col1, col2) in enumerate(yt_data):
    row = table4.rows[i]
    if i == 0:
        shade_cell(row.cells[0], "7030A0")
        shade_cell(row.cells[1], "7030A0")
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5, bold=True, color=(255, 255, 255))
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5, bold=True, color=(255, 255, 255))
    else:
        fill = "EAD1F5" if i % 2 == 1 else "F9F0FF"
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], fill)
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5, bold=True if i == 1 else False)
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5)

for row in table4.rows:
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(10.5)

doc.add_paragraph()

# ===== 週間スケジュール =====
add_heading(doc, "■ 週間スケジュール例")

table5 = doc.add_table(rows=3, cols=2)
table5.style = 'Table Grid'
schedule = [
    ("区分", "内容"),
    ("平日（月〜金）\n1時間", "・30分：テキスト or 問題集\n・30分：過去問アプリ（通勤・隙間時間）"),
    ("休日（土・日）\n2〜3時間", "・1時間：前週の復習\n・1〜2時間：新単元 or 模擬試験"),
]

for i, (col1, col2) in enumerate(schedule):
    row = table5.rows[i]
    if i == 0:
        shade_cell(row.cells[0], "1F497D")
        shade_cell(row.cells[1], "1F497D")
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5, bold=True, color=(255, 255, 255))
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5, bold=True, color=(255, 255, 255))
    else:
        fill = "DCE6F1" if i % 2 == 1 else "EBF3FB"
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], fill)
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        set_font(r0, size=10.5, bold=True)
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        set_font(r1, size=10.5)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row in table5.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

doc.add_paragraph()

# ===== アドバイス =====
add_heading(doc, "■ 合格のためのアドバイス")

advice_items = [
    "テキスト学習が主、YouTubeは補助として使うのがベスト",
    "理解が難しい単元（年金・相続税の計算）はYouTubeで視覚的に理解する",
    "タックスと相続は数字の暗記が多いため、早めに取り組む",
    "倍速（1.5〜2倍速）でYouTubeを視聴すると時間効率が上がる",
    "視聴後すぐに問題集で確認する習慣をつける",
    "まずはTAC出版「みんなが欲しかった！FPの教科書3級」から始めるのがおすすめ",
]

for item in advice_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    set_font(run, size=10.5)

# 保存
output_path = r"c:\.github\zumi sysytem\FP3級合格ロードマップ.docx"
doc.save(output_path)
print(f"保存完了: {output_path}")
